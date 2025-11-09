import logging
from typing import List, Dict
from pathlib import Path
import tempfile
import io

from PyPDF2 import PdfReader, PdfWriter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup

from app.models import Translation, TranslationChunk

logger = logging.getLogger(__name__)


class BookAssembler:
    """Monta livros traduzidos nos formatos PDF, EPUB e DOCX"""
    
    def __init__(self, translation: Translation, chunks: List[TranslationChunk]):
        self.translation = translation
        self.chunks = sorted(chunks, key=lambda x: x.chunk_order)
        logger.info(f"Inicializando BookAssembler para tradução {translation.id} com {len(chunks)} chunks")
    
    def assemble_book(self, output_path: Path) -> Path:
        """Monta o livro no formato especificado"""
        logger.info(f"Montando livro traduzido em: {output_path}")
        
        try:
            # Garantir que o diretório existe
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Chamar método apropriado baseado no formato
            format_methods = {
                'pdf': self.assemble_pdf,
                'epub': self.assemble_epub,
                'docx': self.assemble_docx,
                'doc': self.assemble_docx
            }
            
            file_format = self.translation.file_format.lower()
            if file_format not in format_methods:
                raise ValueError(f"Formato não suportado: {file_format}")
            
            return format_methods[file_format](output_path)
            
        except Exception as e:
            logger.error(f"Erro ao montar livro: {str(e)}")
            raise
    
    def assemble_pdf(self, output_path: Path) -> Path:
        """Monta PDF traduzido usando reportlab"""
        logger.info("Montando PDF traduzido...")
        
        try:
            # Criar PDF temporário com reportlab
            packet = io.BytesIO()
            can = canvas.Canvas(packet, pagesize=letter)
            
            # Configurações de página
            width, height = letter
            margin = 72  # 1 inch
            line_height = 14
            y_position = height - margin
            
            # Adicionar metadados
            can.setTitle(f"{self.translation.title} - Traduzido")
            can.setAuthor("BookAI Translation System")
            
            # Processar cada chunk
            for i, chunk in enumerate(self.chunks):
                if i > 0 and y_position < margin + (line_height * 5):
                    can.showPage()
                    y_position = height - margin
                
                # Adicionar texto do chunk
                text = chunk.translated_text or chunk.original_text
                lines = self._wrap_text(text, can, width - (2 * margin))
                
                for line in lines:
                    if y_position < margin + line_height:
                        can.showPage()
                        y_position = height - margin
                    
                    can.drawString(margin, y_position, line)
                    y_position -= line_height
                
                # Espaço entre chunks
                y_position -= line_height
            
            can.save()
            
            # Mover para posição inicial do buffer
            packet.seek(0)
            
            # Salvar arquivo final
            with open(output_path, 'wb') as f:
                f.write(packet.getvalue())
            
            logger.info(f"PDF montado com sucesso: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Erro ao montar PDF: {str(e)}")
            raise
    
    def assemble_docx(self, output_path: Path) -> Path:
        """Monta documento DOCX traduzido"""
        logger.info("Montando DOCX traduzido...")
        
        try:
            doc = Document()
            
            # Adicionar título
            title = doc.add_heading(f"{self.translation.title} - Traduzido", 0)
            
            # Processar cada chunk
            for chunk in self.chunks:
                text = chunk.translated_text or chunk.original_text
                
                # Adicionar parágrafo com o texto traduzido
                paragraph = doc.add_paragraph(text)
                
                # Adicionar quebra entre chunks (exceto no último)
                if chunk != self.chunks[-1]:
                    doc.add_paragraph()  # parágrafo vazio como separador
            
            # Salvar documento
            doc.save(output_path)
            logger.info(f"DOCX montado com sucesso: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Erro ao montar DOCX: {str(e)}")
            raise
    
    def assemble_epub(self, output_path: Path) -> Path:
        """Monta EPUB traduzido"""
        logger.info("Montando EPUB traduzido...")
        
        try:
            # Criar novo livro EPUB
            book = epub.EpubBook()
            
            # Metadados
            book.set_identifier(f"translation_{self.translation.id}")
            book.set_title(f"{self.translation.title} - Traduzido")
            book.set_language('pt')  # Português como língua da tradução
            book.add_author("BookAI Translation System")
            
            # Criar capítulos para cada chunk
            chapters = []
            spine = ['nav']
            
            for i, chunk in enumerate(self.chunks):
                text = chunk.translated_text or chunk.original_text
                
                # Criar capítulo
                chapter = epub.EpubHtml(
                    title=f'Capítulo {i+1}',
                    file_name=f'chap_{i+1}.xhtml',
                    lang='pt'
                )
                
                # Conteúdo do capítulo
                chapter.content = f'''
                    <html>
                    <head>
                        <title>Capítulo {i+1}</title>
                    </head>
                    <body>
                        <h1>Capítulo {i+1}</h1>
                        <p>{text}</p>
                    </body>
                    </html>
                '''
                
                # Adicionar capítulo ao livro
                book.add_item(chapter)
                chapters.append(chapter)
                spine.append(chapter)
            
            # Definir estrutura do livro
            book.toc = chapters
            book.spine = spine
            
            # Adicionar navegação padrão
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            
            # Salvar EPUB
            epub.write_epub(output_path, book, {})
            logger.info(f"EPUB montado com sucesso: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Erro ao montar EPUB: {str(e)}")
            raise
    
    def _wrap_text(self, text: str, canvas: canvas.Canvas, max_width: float) -> List[str]:
        """Quebra texto em linhas que cabem na largura especificada"""
        lines = []
        words = text.split()
        
        current_line = []
        current_width = 0
        
        for word in words:
            word_width = canvas.stringWidth(word + ' ', 'Helvetica', 12)
            
            if current_width + word_width <= max_width:
                current_line.append(word)
                current_width += word_width
            else:
                if current_line:
                    lines.append(' '.join(current_line))
                current_line = [word]
                current_width = word_width
        
        if current_line:
            lines.append(' '.join(current_line))
        
        return lines


def assemble_translated_book(
    translation: Translation,
    chunks: List[TranslationChunk],
    output_dir: Path
) -> Path:
    """
    Helper: montar livro traduzido
    
    Args:
        translation: Objeto Translation
        chunks: Lista de TranslationChunk ordenados
        output_dir: Diretório de saída
    
    Returns:
        Path do arquivo gerado
    """
    logger.info(f"Montando livro traduzido para tradução {translation.id}")
    
    try:
        assembler = BookAssembler(translation, chunks)
        
        # Gerar nome do arquivo
        filename = f"{translation.id}_translated.{translation.file_format}"
        output_path = output_dir / filename
        
        return assembler.assemble_book(output_path)
        
    except Exception as e:
        logger.error(f"Erro na função assemble_translated_book: {str(e)}")
        raise
