import logging
from docx import Document
from typing import Optional, Dict, Any
import os

logger = logging.getLogger(__name__)


class DOCXParser:
    """Parser para arquivos DOCX"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = None
        self._load_docx()
    
    def _load_docx(self):
        """Carregar DOCX"""
        try:
            self.doc = Document(self.file_path)
            logger.info(f"DOCX carregado: {self.file_path}")
        except Exception as e:
            logger.error(f"Erro ao carregar DOCX: {str(e)}")
            raise ValueError(f"Erro ao carregar DOCX: {str(e)}")
    
    def extract_text(self, paragraph_range: Optional[tuple] = None) -> str:
        """Extrair texto do DOCX"""
        try:
            paragraphs = self.doc.paragraphs
            
            if paragraph_range:
                start, end = paragraph_range
                paragraphs = paragraphs[start:end]
            
            texts = [para.text for para in paragraphs if para.text.strip()]
            full_text = "\n\n".join(texts)
            
            logger.debug(f"Extraídos {len(full_text)} caracteres de {len(texts)} parágrafos")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Erro ao extrair texto do DOCX: {str(e)}")
            raise ValueError(f"Erro ao extrair texto: {str(e)}")
    
    def extract_metadata(self) -> Dict[str, Any]:
        """Extrair metadados do DOCX"""
        try:
            core_properties = self.doc.core_properties
            
            return {
                "title": core_properties.title or "Sem título",
                "author": core_properties.author or "Desconhecido",
                "paragraphs": len(self.doc.paragraphs),
                "created": core_properties.created,
                "modified": core_properties.modified
            }
        except Exception as e:
            logger.error(f"Erro ao extrair metadados: {str(e)}")
            return {
                "title": "Sem título",
                "author": "Desconhecido",
                "paragraphs": len(self.doc.paragraphs) if self.doc else 0,
                "created": None,
                "modified": None
            }
    
    def get_total_paragraphs(self) -> int:
        """Retornar total de parágrafos não vazios"""
        return len([p for p in self.doc.paragraphs if p.text.strip()])
    
    def calculate_characters(self, paragraph_range: Optional[tuple] = None) -> int:
        """Calcular total de caracteres"""
        text = self.extract_text(paragraph_range)
        return len(text)
    
    def estimate_pages(self) -> int:
        """Estimar número de páginas (baseado em 2500 chars/página)"""
        total_chars = self.calculate_characters()
        return max(1, total_chars // 2500)


# Funções helper
def extract_text_from_docx(file_path: str) -> str:
    """Helper: extrair texto completo"""
    parser = DOCXParser(file_path)
    return parser.extract_text()


def extract_preview_from_docx(file_path: str, max_paragraphs: int = 50) -> str:
    """Helper: extrair preview de N parágrafos"""
    parser = DOCXParser(file_path)
    return parser.extract_text((0, max_paragraphs))


def get_docx_info(file_path: str) -> Dict[str, Any]:
    """Helper: obter informações do DOCX"""
    try:
        parser = DOCXParser(file_path)
        metadata = parser.extract_metadata()
        total_chars = parser.calculate_characters()
        
        return {
            **metadata,
            "total_characters": total_chars,
            "estimated_pages": parser.estimate_pages(),
            "file_size_mb": round(os.path.getsize(file_path) / 1024 / 1024, 2)
        }
    except Exception as e:
        logger.error(f"Erro ao obter info do DOCX: {str(e)}")
        raise
